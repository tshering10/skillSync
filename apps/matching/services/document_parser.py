import re
from pathlib import Path
import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from sentence_transformers import SentenceTransformer
import spacy

# NLP / Model loading
try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    nlp = None

model = None

def get_sentence_transformer():
    """Lazily load and return the SentenceTransformer model."""
    global model
    if model is None:
        model = SentenceTransformer("all-MiniLM-L6-v2")
    return model


# Text cleaning & normalization
def clean_extracted_text(text: str) -> str:
    """Normalize whitespace, unicode ligatures, bullets, and linebreaks."""
    if not text:
        return ""

    # Replace common unicode ligatures & special spaces
    replacements = {
        "\xa0": " ",
        "\u200b": " ",
        "\ufb01": "fi",
        "\ufb02": "fl",
        "\ufb00": "ff",
        "\ufb03": "ffi",
        "\ufb04": "ffl",
        "\r\n": "\n",
        "\r": "\n",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # Normalize bullet points into clean dashes
    text = re.sub(r"[\u2022\u2023\u25E6\u2043\u2219\u25AA\u25AB\u25CF\u25CB\uf0b7\uf0a7]", "\n- ", text)

    # Fix broken hyphenated line breaks (e.g. "experi-\nence" -> "experience")
    text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)

    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Strip whitespace per line
    lines = [line.strip() for line in text.split("\n")]
    return "\n".join(lines).strip()


# PDF extraction
def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using PyMuPDF block layout extraction with pdfplumber fallback.
    Maintains column reading order for multi-column resumes.
    """
    extracted_text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            # Extract text blocks: (x0, y0, x1, y1, text, block_no, block_type)
            blocks = page.get_text("blocks")
            # Sort primarily top-to-bottom, secondarily left-to-right
            blocks = sorted(blocks, key=lambda b: (round(b[1] / 20), b[0]))
            for b in blocks:
                block_text = b[4].strip()
                if block_text:
                    extracted_text += block_text + "\n"
        doc.close()
    except Exception:
        pass

    # Fallback to pdfplumber if fitz produced little/no text
    if not extracted_text.strip():
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {str(e)}")

    if not extracted_text.strip():
        raise ValueError("Could not extract text from PDF. Scanned/image-only PDFs require OCR.")

    return clean_extracted_text(extracted_text)


# DOCX extraction
def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX paragraphs and tables."""
    try:
        doc = Document(file_path)
        full_text = []

        # 1. Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. Tables (captures tabular skills/experience sections)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    # Remove duplicate cells created by merged columns
                    unique_cells = list(dict.fromkeys(row_cells))
                    full_text.append(" | ".join(unique_cells))

        return clean_extracted_text("\n".join(full_text))
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")


# Plain text extraction
def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text file with fallback encoding."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return clean_extracted_text(f.read())
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            return clean_extracted_text(f.read())
    except Exception as e:
        raise ValueError(f"Failed to parse text file: {str(e)}")


# Universal document entrypoint
def parse_document(file_path: str) -> str:
    """Detect file extension and extract cleaned text."""
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension == ".docx":
        return extract_text_from_docx(file_path)
    elif extension == ".txt":
        return extract_text_from_txt(file_path)
    elif extension == ".doc":
        raise ValueError("Legacy '.doc' format is not supported. Please convert to modern '.docx' or '.pdf'.")
    else:
        raise ValueError(f"Unsupported file format '{extension}'. Please provide a PDF, DOCX, or TXT file.")


# Document semantic encoding
def encode_document(text: str, transformer=None, chunk_size: int = 1500, overlap: int = 250):
    """Encode complete document with chunked mean-pooling."""
    if transformer is None:
        transformer = get_sentence_transformer()

    if not text:
        return transformer.encode("", convert_to_tensor=True)

    text = text.strip()
    if len(text) <= chunk_size:
        return transformer.encode(text, convert_to_tensor=True)

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start += chunk_size - overlap

    chunk_embeddings = transformer.encode(chunks, convert_to_tensor=True)
    mean_embedding = chunk_embeddings.mean(dim=0, keepdim=True)
    normalized_embedding = mean_embedding / mean_embedding.norm(dim=1, keepdim=True)
    return normalized_embedding.squeeze(0)
